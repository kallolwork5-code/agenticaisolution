import os
from typing import Dict, List, Optional
from pathlib import Path
import hashlib
import uuid
from datetime import datetime
import logging

# Document processing imports
import PyPDF2
from docx import Document
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer

# Text processing
import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

logger = logging.getLogger(__name__)

class DocumentProcessingAgent:
    """
    Agent responsible for processing PDF and Word documents for RAG
    - Extracts text content
    - Chunks and tokenizes text
    - Creates embeddings
    - Stores in ChromaDB
    """
    
    def __init__(self, 
                 chroma_db_path: str = "data/chroma_db",
                 embedding_model: str = "all-MiniLM-L6-v2"):
        self.chroma_db_path = chroma_db_path
        self.embedding_model_name = embedding_model
        
        # Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(
            path=chroma_db_path,
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Initialize embedding model
        self.embedding_model = SentenceTransformer(embedding_model)
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Ensure directory exists
        Path(chroma_db_path).mkdir(parents=True, exist_ok=True)
    
    def process_document(self, 
                        file_path: str, 
                        file_name: str, 
                        file_type: str,
                        file_id: str,
                        content_hash: str) -> Dict:
        """
        Main method to process documents for RAG
        """
        try:
            # Check for duplicates
            duplicate_info = self._check_duplicate(content_hash)
            if duplicate_info['is_duplicate']:
                return {
                    'status': 'duplicate',
                    'message': f'Document is a duplicate of {duplicate_info["original_file"]}',
                    'chunks_processed': 0,
                    'duplicate_info': duplicate_info
                }
            
            # Extract text content
            text_content = self._extract_text(file_path, file_type)
            if not text_content.strip():
                return {
                    'status': 'error',
                    'message': 'No text content could be extracted from the document',
                    'chunks_processed': 0
                }
            
            # Create document chunks
            chunks = self._create_chunks(text_content, file_name, file_id)
            
            # Generate embeddings and store in ChromaDB
            collection_name = self._get_collection_name(file_type)
            chunks_stored = self._store_in_chromadb(chunks, collection_name, file_id, content_hash)
            
            # Store metadata
            self._store_document_metadata(file_id, file_name, content_hash, len(chunks), collection_name)
            
            return {
                'status': 'success',
                'message': f'Successfully processed document into {len(chunks)} chunks',
                'chunks_processed': len(chunks),
                'collection_name': collection_name,
                'text_length': len(text_content),
                'chunks_stored': chunks_stored
            }
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                'status': 'error',
                'message': f'Document processing failed: {str(e)}',
                'chunks_processed': 0
            }
    
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from different document types"""
        try:
            if file_type == 'application/pdf':
                return self._extract_pdf_text(file_path)
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword']:
                return self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported document type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error reading Word document: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-.,;:!?()[\]{}"\'/\\]', ' ', text)
        
        return text.strip()
    
    def _create_chunks(self, text_content: str, file_name: str, file_id: str) -> List[Dict]:
        """Create text chunks for embedding"""
        # Create LangChain document
        doc = LangChainDocument(
            page_content=text_content,
            metadata={
                'source': file_name,
                'file_id': file_id,
                'processed_at': datetime.utcnow().isoformat()
            }
        )
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Convert to our format
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            processed_chunks.append({
                'chunk_id': chunk_id,
                'content': chunk.page_content,
                'metadata': {
                    **chunk.metadata,
                    'chunk_index': i,
                    'chunk_length': len(chunk.page_content)
                }
            })
        
        return processed_chunks
    
    def _get_collection_name(self, file_type: str) -> str:
        """Get ChromaDB collection name based on file type"""
        if file_type == 'application/pdf':
            return 'pdf_documents'
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          'application/msword']:
            return 'word_documents'
        else:
            return 'documents'
    
    def _store_in_chromadb(self, chunks: List[Dict], collection_name: str, file_id: str, content_hash: str) -> int:
        """Store chunks in ChromaDB with embeddings"""
        try:
            # Get or create collection
            collection = self.chroma_client.get_or_create_collection(
                name=collection_name,
                metadata={"description": f"Collection for {collection_name}"}
            )
            
            # Prepare data for ChromaDB
            documents = []
            metadatas = []
            ids = []
            
            for chunk in chunks:
                documents.append(chunk['content'])
                metadatas.append({
                    **chunk['metadata'],
                    'content_hash': content_hash
                })
                ids.append(chunk['chunk_id'])
            
            # Generate embeddings
            embeddings = self.embedding_model.encode(documents).tolist()
            
            # Store in ChromaDB
            collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids,
                embeddings=embeddings
            )
            
            logger.info(f"Stored {len(chunks)} chunks in collection {collection_name}")
            return len(chunks)
            
        except Exception as e:
            logger.error(f"Error storing in ChromaDB: {str(e)}")
            raise
    
    def _check_duplicate(self, content_hash: str) -> Dict:
        """Check if document content already exists in any collection"""
        try:
            collections = self.chroma_client.list_collections()
            
            for collection in collections:
                coll = self.chroma_client.get_collection(collection.name)
                results = coll.get(
                    where={"content_hash": content_hash},
                    limit=1
                )
                
                if results['ids']:
                    return {
                        'is_duplicate': True,
                        'original_file': results['metadatas'][0].get('source', 'Unknown'),
                        'collection': collection.name,
                        'processed_at': results['metadatas'][0].get('processed_at')
                    }
            
            return {'is_duplicate': False}
            
        except Exception as e:
            logger.error(f"Error checking duplicates: {str(e)}")
            return {'is_duplicate': False}
    
    def _store_document_metadata(self, file_id: str, file_name: str, content_hash: str, chunk_count: int, collection_name: str):
        """Store document processing metadata"""
        # This could be stored in SQLite or another metadata store
        # For now, we'll log it
        logger.info(f"Document processed: {file_name} -> {chunk_count} chunks in {collection_name}")
    
    def search_documents(self, query: str, collection_name: Optional[str] = None, limit: int = 5) -> List[Dict]:
        """Search documents using semantic similarity"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query]).tolist()[0]
            
            results = []
            collections_to_search = [collection_name] if collection_name else [
                'pdf_documents', 'word_documents', 'documents'
            ]
            
            for coll_name in collections_to_search:
                try:
                    collection = self.chroma_client.get_collection(coll_name)
                    search_results = collection.query(
                        query_embeddings=[query_embedding],
                        n_results=limit
                    )
                    
                    for i, doc_id in enumerate(search_results['ids'][0]):
                        results.append({
                            'id': doc_id,
                            'content': search_results['documents'][0][i],
                            'metadata': search_results['metadatas'][0][i],
                            'distance': search_results['distances'][0][i],
                            'collection': coll_name
                        })
                        
                except Exception as e:
                    logger.warning(f"Error searching collection {coll_name}: {str(e)}")
                    continue
            
            # Sort by distance (similarity)
            results.sort(key=lambda x: x['distance'])
            return results[:limit]
            
        except Exception as e:
            logger.error(f"Error searching documents: {str(e)}")
            return []
    
    def get_document_info(self, file_id: str) -> Dict:
        """Get information about a processed document"""
        try:
            collections = self.chroma_client.list_collections()
            
            for collection in collections:
                coll = self.chroma_client.get_collection(collection.name)
                results = coll.get(
                    where={"file_id": file_id}
                )
                
                if results['ids']:
                    return {
                        'file_id': file_id,
                        'collection': collection.name,
                        'chunk_count': len(results['ids']),
                        'chunks': [
                            {
                                'id': results['ids'][i],
                                'content_preview': results['documents'][i][:200] + '...',
                                'metadata': results['metadatas'][i]
                            }
                            for i in range(len(results['ids']))
                        ]
                    }
            
            return {'error': 'Document not found'}
            
        except Exception as e:
            logger.error(f"Error getting document info: {str(e)}")
            return {'error': str(e)}