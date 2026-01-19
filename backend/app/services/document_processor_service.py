"""
Enhanced Document Processing Service

Handles PDF, Word, and other document formats with:
- Text extraction and chunking
- Vector storage in ChromaDB
- Document summarization using AI
- Metadata extraction
"""

import os
import io
import json
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import logging

# Document processing libraries
import pandas as pd
from docx import Document
from pypdf import PdfReader
import fitz  # PyMuPDF for better PDF processing

# Text processing
import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

# Vector storage
from app.vectorstore.chroma_client import ChromaClient
from app.services.ai_insights_service import ai_insights_service

logger = logging.getLogger("document-processor")

class DocumentProcessor:
    def __init__(self):
        self.chroma_client = ChromaClient()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    async def process_document(self, file_content: bytes, filename: str, file_type: str, websocket_manager=None) -> Dict[str, Any]:
        """
        Process a document file with chunking and vector storage
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            file_type: File extension (pdf, docx, doc, txt)
            websocket_manager: Optional WebSocket manager for progress updates
            
        Returns:
            Dict with processing results including chunks, summary, and metadata
        """
        
        try:
            # Extract text content based on file type
            if websocket_manager:
                await websocket_manager.send_progress_update(
                    "analyze", 20, f"Extracting text from {file_type.upper()} document"
                )
            
            text_content, metadata = await self._extract_text_content(
                file_content, filename, file_type
            )
            
            if websocket_manager:
                await websocket_manager.send_agent_update(
                    "analysis",
                    f"Extracted {len(text_content)} characters from document",
                    0.9,
                    {"text_length": len(text_content), "file_type": file_type}
                )
            
            # Chunk the document
            if websocket_manager:
                await websocket_manager.send_progress_update(
                    "analyze", 40, "Chunking document for optimal processing"
                )
            
            chunks = await self._chunk_document(text_content, metadata)
            
            if websocket_manager:
                await websocket_manager.send_agent_update(
                    "analysis",
                    f"Created {len(chunks)} chunks for vector storage",
                    0.85,
                    {"chunk_count": len(chunks), "avg_chunk_size": sum(len(c.page_content) for c in chunks) // len(chunks)}
                )
            
            # Generate document summary using document_summarizer prompts
            if websocket_manager:
                await websocket_manager.send_progress_update(
                    "classify", 60, "Generating AI summary using Document Summarizer"
                )
            
            summary = await self._generate_document_summary(text_content, filename, websocket_manager)
            
            # Store chunks in vector database
            if websocket_manager:
                await websocket_manager.send_progress_update(
                    "store", 80, "Storing document chunks in vector database"
                )
            
            document_id = await self._store_document_chunks(chunks, filename, metadata, summary)
            
            if websocket_manager:
                await websocket_manager.send_agent_update(
                    "decision",
                    f"Document successfully processed and stored with ID: {document_id}",
                    0.95,
                    {"document_id": document_id, "storage_location": "chromadb"}
                )
            
            return {
                "document_id": document_id,
                "text_length": len(text_content),
                "chunk_count": len(chunks),
                "summary": summary,
                "metadata": metadata,
                "storage_location": "chromadb",
                "classification": "document"
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            if websocket_manager:
                await websocket_manager.send_agent_update(
                    "analysis",
                    f"Document processing failed: {str(e)}",
                    0.1,
                    {"error": str(e)}
                )
            raise
    
    async def _extract_text_content(self, file_content: bytes, filename: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text content and metadata from various document formats"""
        
        metadata = {
            "filename": filename,
            "file_type": file_type,
            "processed_at": datetime.now().isoformat(),
            "file_size": len(file_content)
        }
        
        try:
            if file_type.lower() == 'pdf':
                return await self._extract_pdf_content(file_content, metadata)
            elif file_type.lower() in ['docx', 'doc']:
                return await self._extract_word_content(file_content, metadata)
            elif file_type.lower() == 'txt':
                return await self._extract_text_content_plain(file_content, metadata)
            else:
                # Fallback: try to decode as text
                text = file_content.decode('utf-8', errors='ignore')
                metadata["extraction_method"] = "fallback_text"
                return text, metadata
                
        except Exception as e:
            logger.error(f"Error extracting content from {filename}: {e}")
            # Fallback to basic text extraction
            text = file_content.decode('utf-8', errors='ignore')
            metadata["extraction_method"] = "fallback_error"
            metadata["extraction_error"] = str(e)
            return text, metadata
    
    async def _extract_pdf_content(self, file_content: bytes, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from PDF using PyMuPDF for better quality"""
        
        try:
            # Use PyMuPDF for better text extraction
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            text_parts = []
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            
            metadata.update({
                "page_count": page_count,
                "extraction_method": "pymupdf",
                "has_text": len(full_text.strip()) > 0
            })
            
            return full_text, metadata
            
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed, falling back to pypdf: {e}")
            
            # Fallback to pypdf
            try:
                pdf_file = io.BytesIO(file_content)
                reader = PdfReader(pdf_file)
                
                text_parts = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {i + 1} ---\n{text}")
                
                full_text = "\n\n".join(text_parts)
                
                metadata.update({
                    "page_count": len(reader.pages),
                    "extraction_method": "pypdf_fallback",
                    "has_text": len(full_text.strip()) > 0
                })
                
                return full_text, metadata
                
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {e2}")
                metadata["extraction_error"] = f"PyMuPDF: {e}, PyPDF: {e2}"
                return "", metadata
    
    async def _extract_word_content(self, file_content: bytes, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from Word documents"""
        
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine all text
            full_text_parts = []
            if paragraphs:
                full_text_parts.append("\n\n".join(paragraphs))
            if tables_text:
                full_text_parts.append("\n\n--- TABLES ---\n\n" + "\n\n".join(tables_text))
            
            full_text = "\n\n".join(full_text_parts)
            
            metadata.update({
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "extraction_method": "python_docx",
                "has_tables": len(doc.tables) > 0
            })
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Word document extraction failed: {e}")
            metadata["extraction_error"] = str(e)
            return "", metadata
    
    async def _extract_text_content_plain(self, file_content: bytes, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from plain text files"""
        
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    metadata.update({
                        "encoding": encoding,
                        "extraction_method": "text_decode"
                    })
                    return text, metadata
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use error handling
            text = file_content.decode('utf-8', errors='ignore')
            metadata.update({
                "encoding": "utf-8_with_errors",
                "extraction_method": "text_decode_fallback"
            })
            return text, metadata
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            metadata["extraction_error"] = str(e)
            return "", metadata
    
    async def _chunk_document(self, text: str, metadata: Dict) -> List[LangChainDocument]:
        """Chunk document text for vector storage"""
        
        if not text.strip():
            return []
        
        # Clean and normalize text
        cleaned_text = self._clean_text(text)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(cleaned_text)
        
        # Create LangChain documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only include non-empty chunks
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_index": i,
                    "chunk_id": f"{metadata['filename']}_{i}",
                    "chunk_size": len(chunk)
                })
                
                documents.append(LangChainDocument(
                    page_content=chunk,
                    metadata=chunk_metadata
                ))
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        
        # Remove page headers/footers patterns (common in PDFs)
        text = re.sub(r'--- Page \d+ ---\n?', '', text)
        
        # Clean up common PDF artifacts
        text = re.sub(r'\x0c', '\n', text)  # Form feed to newline
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Non-ASCII to space
        
        return text.strip()
    
    async def _generate_document_summary(self, text: str, filename: str, websocket_manager=None) -> str:
        """Generate document summary using the document_summarizer agent"""
        
        try:
            if websocket_manager:
                await websocket_manager.send_agent_update(
                    "prompt",
                    "Using Document Summarizer prompts for comprehensive analysis",
                    0.9,
                    {"agent_role": "document_summarizer"}
                )
            
            # Use the document_summarizer prompts to generate summary
            summary = await ai_insights_service.generate_document_summary(
                text, filename, websocket_manager
            )
            
            return summary
            
        except Exception as e:
            logger.error(f"Error generating document summary: {e}")
            # Fallback to basic summary
            preview = text[:1000] + "..." if len(text) > 1000 else text
            return f"Document: {filename}\n\nContent Preview:\n{preview}"
    
    async def _store_document_chunks(self, chunks: List[LangChainDocument], filename: str, 
                                   metadata: Dict, summary: str) -> str:
        """Store document chunks in ChromaDB vector database"""
        
        try:
            # Generate unique document ID
            document_id = f"doc_{hashlib.md5(filename.encode()).hexdigest()}_{int(datetime.now().timestamp())}"
            
            # Add document-level metadata to all chunks
            for chunk in chunks:
                chunk.metadata.update({
                    "document_id": document_id,
                    "document_summary": summary[:500],  # Truncated summary
                    "total_chunks": len(chunks)
                })
            
            # Store in ChromaDB
            collection_name = "documents"
            await self.chroma_client.add_documents(chunks, collection_name)
            
            logger.info(f"Stored {len(chunks)} chunks for document {document_id}")
            
            return document_id
            
        except Exception as e:
            logger.error(f"Error storing document chunks: {e}")
            raise

# Global instance
document_processor = DocumentProcessor()